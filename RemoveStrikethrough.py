from docx import Document
import re
import os

def removerTextoTachado(caminho_entrada, caminho_saida):
    doc = Document(caminho_entrada)

    # Processa parágrafos
    processar_paragrafos(doc.paragraphs)

    # Processa tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                processar_paragrafos(celula.paragraphs)

    # Processa cabeçalhos e rodapés
    for secao in doc.sections:
        processar_paragrafos(secao.header.paragraphs)
        processar_paragrafos(secao.footer.paragraphs)

    # Corrige os números ordinais
    alterarSimboloOrdinal(doc)

    # Salva o documento
    doc.save(caminho_saida)
    print(f'Documento salvo em: {caminho_saida}')

    # Abre o arquivo automaticamente
    os.startfile(caminho_saida)


def processar_paragrafos(paragrafos):
    for paragrafo in paragrafos:
        runs_para_remover = []
        for run in paragrafo.runs:
            if run.font.strike is not None and run.font.strike:
                runs_para_remover.append(run)
        for run in runs_para_remover:
            run.text = ''  # Remove texto tachado


def alterarSimboloOrdinal(doc):
    padrao_ordinais = r'(\d+)(o|a|s)(-[A-Za-z])?(?=\s|$)'

    # Processa parágrafos normais
    substituir_ordinais(doc.paragraphs)

    # Processa tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_ordinais(celula.paragraphs)

    # Processa cabeçalhos e rodapés
    for secao in doc.sections:
        substituir_ordinais(secao.header.paragraphs)
        substituir_ordinais(secao.footer.paragraphs)


def substituir_ordinais(paragrafos):
    padrao_ordinais = r'(\d+)(o|a|s)(-[A-Za-z])?(?=\s|$)'

    for paragrafo in paragrafos:
        for run in paragrafo.runs:
            if run.font.strike:
                continue  # Ignora texto tachado
            texto_corrigido = re.sub(padrao_ordinais, r'\1°\3', run.text)
            run.text = texto_corrigido


# Exemplo de uso - substitua pelos seus caminhos
removerTextoTachado(
    r'C:\Users\edito\OneDrive\Matheus Siqueira\Python\RemoveStrikethrough\MINI CÓDIGO CIVIL - (23.06).docx',
    r'C:\Users\edito\OneDrive\Matheus Siqueira\Python\RemoveStrikethrough\saida_sem_tachado.docx'
)